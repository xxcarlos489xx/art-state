# scripts/consulta_vector_db.py
import sys
import os
from datetime import datetime
import traceback
import mysql.connector

from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain.vectorstores import Chroma
from langchain.prompts import ChatPromptTemplate
from langchain.schema.runnable import RunnableMap
from langchain.schema.output_parser import StrOutputParser
from dotenv import load_dotenv
from docx import Document
import re
from fpdf import FPDF

# --- CONFIGURACIÓN Y LOGS (Consistente con tus otros scripts) ---
script_dir  =   os.path.dirname(os.path.abspath(__file__))
dotenv_path =   os.path.join(script_dir, "..", ".env")
log_path    =   os.path.join(script_dir, "log_sota.txt")
load_dotenv(dotenv_path)

def write_log(message):
    """Función para escribir en el log con timestamp."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(f"{timestamp} - {message}\n")

def format_docs_with_sources(docs):
    """
    Formatea los documentos recuperados para incluir su contenido y metadatos
    de una manera clara para el LLM.
    """
    formatted_docs = []
    for i, doc in enumerate(docs):
        # Usamos .get() para evitar errores si un metadato no existe.
        paper_id = doc.metadata.get('paper_id', 'ID Desconocido')
        source_file = doc.metadata.get('source', 'Fuente Desconocida')
        
        # Creamos una entrada formateada para cada documento.
        doc_string = (
            f"Fuente [{i+1}]:\n"
            f"  Paper ID: {paper_id}\n"
            f"  Documento: {os.path.basename(source_file)}\n"
            f"  Contenido: \"{doc.page_content}\""
        )
        formatted_docs.append(doc_string)
        
    # Unimos todas las entradas formateadas en un solo bloque de texto.
    return "\n\n---\n\n".join(formatted_docs)

def prepare_context_and_bibliography(docs):
    """
    Procesa los documentos recuperados para crear un mapa de referencias único,
    un contexto formateado para el LLM y una cadena de bibliografía final.
    """
    reference_map = {}
    bibliography_entries = {}
    formatted_context_parts = []
    
    current_ref_number = 1
    
    for doc in docs:
        paper_id = doc.metadata.get('paper_id', 'ID Desconocido')
        
        # Si es un paper nuevo, se le asigna un número de referencia
        if paper_id not in reference_map:
            reference_map[paper_id] = current_ref_number
            
            # Crear la entrada para la bibliografía final
            source_file = os.path.basename(doc.metadata.get('source', 'Fuente Desconocida'))
            bibliography_entries[current_ref_number] = f"[{current_ref_number}] ID del Paper: {paper_id}, Documento: {source_file}"
            
            current_ref_number += 1
            
        # Obtener el número de referencia para este chunk
        ref_num_for_chunk = reference_map[paper_id]
        
        # Formatear el contexto para el LLM
        context_part = (
            f"Fragmento de fuente (citar como [{ref_num_for_chunk}]):\n"
            f"\"{doc.page_content}\""
        )
        formatted_context_parts.append(context_part)
        
    # Construir la cadena de contexto final
    final_context = "\n\n---\n\n".join(formatted_context_parts)
    
    # Construir la cadena de bibliografía final, ordenada por número
    bibliography_string = "Referencias\n"
    for i in sorted(bibliography_entries.keys()):
        bibliography_string += bibliography_entries[i] + "\n"
        
    return final_context, bibliography_string

def save_sota_to_docx(sota_text, topic_title, file_path):
    """Guarda el texto del SOTA en un archivo .docx."""
    try:
        doc = Document()
        doc.add_heading(f"Estado del Arte: {topic_title}", level=1)
        
        intro_text = (
            f"Este documento presenta el Estado del Arte generado automáticamente para el tema de investigación "
            f"mencionado. La siguiente sección ha sido sintetizada a partir de las fuentes documentales proporcionadas y es con fines académicos."
        )
        doc.add_paragraph(intro_text)        
        doc.add_paragraph()

        # añadir párrafos separados
        for paragraph_text in sota_text.split('\n'):
            if paragraph_text.strip():  # Evita añadir párrafos vacíos
                doc.add_paragraph(paragraph_text)
        
        doc.save(file_path)
        write_log(f"Archivo Word guardado exitosamente en: {file_path}")

    except Exception as e:
        write_log(f"ERROR al crear el archivo .docx: {str(e)}")
        write_log(traceback.format_exc())

def save_sota_to_db(topic_id, file_path):
    """
    Calcula la nueva versión y guarda el registro del SOTA en la base de datos.
    """
    try:
        db_connection = mysql.connector.connect(
            host=os.getenv("DB_HOST"),
            user=os.getenv("DB_USER"),
            password=os.getenv("DB_PASSWORD"),
            database=os.getenv("DB_NAME")
        )
        if db_connection.is_connected():
            cursor = db_connection.cursor()
            
            write_log(f"Calculando nueva versión para topic_id: {topic_id}")
            query_max_version = "SELECT MAX(version) FROM sotas WHERE topic_id = %s"
            cursor.execute(query_max_version, (topic_id,))
            result = cursor.fetchone()
            
            current_max_version = result[0] if result and result[0] is not None else 0
            new_version = current_max_version + 1
            write_log(f"La nueva versión será: {new_version}")

            insert_query = """
                INSERT INTO sotas (ruta, topic_id, version)
                VALUES (%s, %s, %s)
            """

            relative_path           =   file_path.split(os.sep + 'storage' + os.sep)[1]
            relative_path_for_url   =   relative_path.replace('\\', '/')
            web_path                =   f"/api/storage/{relative_path_for_url}"

            cursor.execute(insert_query, (web_path, topic_id, new_version))
            db_connection.commit()
            write_log(f"Guardando sota docx: {web_path}")
            write_log(f"Registro de SOTA v{new_version} guardado en la BD para topic_id: {topic_id}")
            
    except mysql.connector.Error as e:
        write_log(f"ERROR de MySQL al guardar el SOTA: {e}")
        write_log(traceback.format_exc())
    finally:
        if 'db_connection' in locals() and db_connection.is_connected():
            cursor.close()
            db_connection.close()

def clean_citations_from_text(text):
    """
    Usa expresiones regulares para eliminar las citas en formato [n] del texto.
    """
    # La expresión regular busca un corchete de apertura [, seguido de uno o más dígitos \d+,
    # y luego un corchete de cierre ]. El espacio \s*? busca opcionalmente espacios antes de la cita.
    cleaned_text = re.sub(r'\s*\[[\d\s,]+\]', '', text)
    return cleaned_text

def save_body_to_pdf(body_text, file_path):
    """
    Guarda el texto del cuerpo del SOTA en un archivo PDF.
    """
    try:
        pdf = FPDF()
        pdf.add_page()
        font_path = os.path.join(script_dir, "DejaVuSans.ttf")
        pdf.add_font("DejaVu", "", font_path, uni=True)
        pdf.set_font("DejaVu", size=12)        
        pdf.multi_cell(0, 5, txt=body_text)        
        pdf.output(file_path)
        write_log(f"PDF del cuerpo guardado exitosamente en: {file_path}")

    except Exception as e:
        # FPDF puede dar un error si no encuentra la fuente. Este es un error común.
        if "FPDF error: Can't open file" in str(e):
            write_log("ERROR FPDF: No se encontró la fuente 'DejaVuSans.ttf'.")
            write_log("Asegúrate de haber descargado la fuente y colocarla en el directorio del script o en una ruta conocida.")
        write_log(f"ERROR al crear el archivo PDF: {str(e)}")
        write_log(traceback.format_exc())

# --- LÓGICA PRINCIPAL ---
if __name__ == "__main__":
    if len(sys.argv) < 3:
        write_log("Faltan parámetros: se requiere topic_id y topic_title.")
        sys.exit(1)

    topic_id            =   sys.argv[1]
    topic_title         =   sys.argv[2]
    output_path         =   sys.argv[3] #route save sota
    persist_directory   =   os.path.join(script_dir, "chroma", f"db_{topic_id}")
    
    write_log(f"Iniciando generación de SOTA para topic_id: {topic_id} con título: '{topic_title}'")

    try:
        # 1. Cargar la base de datos vectorial existente
        embeddings  =   GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        vectordb    =   Chroma(
            persist_directory=persist_directory,
            embedding_function=embeddings
        )
        
        # 2. Crear un "retriever" para buscar documentos relevantes
        # k=15 significa que buscará los 15 chunks más relevantes al tema
        retriever = vectordb.as_retriever(search_kwargs={"k": 10})
        # retriever = vectordb.as_retriever(search_kwargs={"k": 20})
        # retriever = vectordb.as_retriever(search_kwargs={"k": 25})


        # 3. Definir el prompt para Gemini
        # Este es el corazón de la operación. Le damos instrucciones muy precisas.
        template = """
        Eres un experto investigador académico especializado en la redacción de revisiones de literatura y estados del arte (State of the Art).
        Tu tarea es sintetizar la información proporcionada para construir una sección de "Estado del Arte" coherente, bien estructurada y académica sobre el siguiente tema de investigación.

        TEMA DE INVESTIGACIÓN: "{question}"

        Utiliza ÚNICAMENTE la siguiente información extraída de documentos académicos como base para tu redacción. No inventes información ni utilices conocimiento externo.

        CONTEXTO (FUENTES):
        {context}

        --- DIRECTRICES DE REDACCIÓN Y ESTILO ---
        1.  **Redacción Fluida:** Escribe un texto continuo, bien estructurado y académico. Comienza con una introducción, desarrolla el análisis en el cuerpo del texto y finaliza con una conclusión. NO uses subtítulos como "Introducción" o "Conclusión".
        2.  **Citación Obligatoria:** Cuando uses información de un fragmento del contexto, DEBES citarlo usando el número que se te indica. Por ejemplo: `(citar como [1])` significa que debes poner `[1]` en el texto.
        3.  **NO GENERES BIBLIOGRAFÍA:** NO añadas una sección de "Referencias" al final. Yo me encargaré de eso. Tu tarea es únicamente redactar el cuerpo del Estado del Arte con sus citas.
        4.  **Tono:** Mantén un tono formal y objetivo. Si encuentras información contradictoria, señálalo.

        --- INSTRUCCIÓN CRÍTICA ---
        Si el CONTEXTO proporcionado es insuficiente, responde únicamente con: "No se encontró información relevante en los documentos proporcionados para generar un Estado del Arte sobre este tema."

        --- COMIENZA LA REDACCIÓN ---
        """
        prompt = ChatPromptTemplate.from_template(template)

        # 4. Configurar el modelo de Gemini
        llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash-lite", temperature=0.4, convert_system_message_to_human=True)
        
        retrieved_docs = retriever.invoke(topic_title)
        if not retrieved_docs:
            write_log("No se recuperaron documentos. Abortando.")
            # Aquí podrías guardar un SOTA vacío o con un mensaje de error
            sota_result = "No se encontró información relevante en los documentos proporcionados para generar un Estado del Arte sobre este tema."
        else:
            formatted_context, bibliography_string = prepare_context_and_bibliography(retrieved_docs)
            
            # 3. Construir la cadena y ejecutarla
            rag_chain = (
                RunnableMap({
                    "context": lambda x: x["context"],
                    "question": lambda x: x["question"]
                })
                | prompt
                | llm
                | StrOutputParser()
            )

            write_log(f"Invocando la cadena RAG con Gemini...")
            sota_body = rag_chain.invoke({
                "context": formatted_context,
                "question": topic_title
            })
            
            # 4. Ensamblar el resultado final
            sota_result = sota_body + "\n\n" + bibliography_string
     
        write_log(f"SOTA generado exitosamente. Guardando en la base de datos...")

        sota_body_cleaned = clean_citations_from_text(sota_body)

        path_word   =   os.path.join(output_path, "sota.docx")
        path_pdf    =   os.path.join(output_path, "sota_body.pdf")

        save_body_to_pdf(sota_body_cleaned, path_pdf)

        save_sota_to_docx(sota_result, topic_title, path_word)
        save_sota_to_db(topic_id, path_word)
        
        write_log("Proceso completado.")

    except Exception as e:
        write_log(f"ERROR CRÍTICO durante la generación de SOTA: {str(e)}")
        write_log(traceback.format_exc())
        # Opcional: podrías actualizar la BD con un estado de error
        # update_topic_status(topic_id, "error")
        sys.exit(1)